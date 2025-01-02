import requests
from bs4 import BeautifulSoup
import re
import time
from docx import Document

# Função para realizar uma requisição HTTP com verificação SSL
def get_page_content(url):
    try:
        response = requests.get(url, verify=True)
        response.raise_for_status()
        return response.text
    except requests.exceptions.RequestException as e:
        print(f"Erro ao acessar o site {url}: {e}")
        return None

# Função para buscar o nome da empresa
def get_company_name(soup):
    company_name = None
    title_tag = soup.find('title')
    if title_tag:
        company_name = title_tag.text.strip()
    if not company_name:
        meta_tag = soup.find('meta', {'name': 'author'})
        if meta_tag:
            company_name = meta_tag.get('content', '').strip()
        if not company_name:
            h1_tag = soup.find('h1')
            if h1_tag:
                company_name = h1_tag.text.strip()
    return company_name if company_name else "Nome não encontrado"

# Função para extrair telefones
def get_phones(soup):
    phones = []
    phone_regex = re.compile(r'\(?\d{2,3}\)?\s?\d{4,5}-?\d{4}')
    for text in soup.stripped_strings:
        if phone_regex.match(text):
            phones.append(text)
    return phones

# Função para extrair e-mails
def get_emails(soup):
    emails = []
    email_regex = re.compile(r'mailto:[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    for link in soup.find_all('a', href=True):
        if email_regex.match(link['href']):
            emails.append(link['href'])
    return emails

# Função para criar e salvar um documento Word com os resultados
def save_to_docx(results):
    # Criando um novo documento
    doc = Document()
    doc.add_heading('Resultados de Pesquisa', 0)

    for result in results:
        doc.add_heading(f"Empresa: {result['company_name']}", level=1)
        doc.add_paragraph(f"Site: {result['company_url']}")
        doc.add_paragraph(f"Telefones encontrados: {', '.join(result['phones']) if result['phones'] else 'Nenhum telefone encontrado'}")
        doc.add_paragraph(f"E-mails encontrados: {', '.join(result['emails']) if result['emails'] else 'Nenhum e-mail encontrado'}")
        doc.add_paragraph("\n" + "-"*50 + "\n")

    # Salvando o documento
    doc.save("resultados_pesquisa.docx")
    print("Arquivo 'resultados_pesquisa.docx' salvo com sucesso!")

# Função para realizar a pesquisa no Google Custom Search API
def search_google(query, api_key, cse_id):
    url = f"https://www.googleapis.com/customsearch/v1?q={query}&key={api_key}&cx={cse_id}"
    response = requests.get(url)
    if response.status_code == 200:
        return response.json()
    else:
        print(f"Erro ao buscar no Google Custom Search: {response.status_code}")
        return None

# Função principal que realiza a extração de dados
def extract_data_from_search(query, api_key, cse_id):
    # Realiza a busca no Google Custom Search
    search_results = search_google(query, api_key, cse_id)
    
    if search_results and "items" in search_results:
        results = []  # Lista para armazenar os resultados extraídos
        
        for item in search_results["items"]:
            company_url = item["link"]
            print(f"Buscando dados para: {query}")
            print(f"Site encontrado: {company_url}")
            
            # Acessar cada site encontrado e tentar extrair as informações
            company_page = get_page_content(company_url)
            if company_page:
                company_soup = BeautifulSoup(company_page, 'html.parser')
                
                company_name = get_company_name(company_soup)
                phones = get_phones(company_soup)
                emails = get_emails(company_soup)

                # Adicionando o resultado ao documento
                result = {
                    'company_name': company_name,
                    'company_url': company_url,
                    'phones': phones,
                    'emails': emails
                }
                results.append(result)
                
            time.sleep(1)  # Atraso para evitar sobrecarga no servidor
        
        if results:
            save_to_docx(results)  # Salvando todos os resultados em um arquivo docx
    else:
        print("Não foi possível obter os resultados de pesquisa.")

if __name__ == "__main__":
    # Defina sua chave de API e ID do mecanismo de pesquisa
    API_KEY = "SUA_CHAVE_DE_API"
    CSE_ID = "SEU_CSE_ID"

    empresas = [
        "Florenzano Nuts",
        "Castanha do Pará Nutriaco",
        "Imaflora",
        "Fábrica de Castanha Benedito Mutran e Cia Ltda",
        "Cooperativa Comaru"
        # ... adicione mais empresas conforme necessário
    ]
    
    for empresa in empresas:
        extract_data_from_search(empresa, API_KEY, CSE_ID)
