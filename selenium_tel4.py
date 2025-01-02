from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.keys import Keys
import time
import re
from bs4 import BeautifulSoup
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document

# Função para inicializar o Selenium
def get_driver():
    driver_service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=driver_service)
    return driver

# Função para extrair telefones de uma página
def get_phones_from_contact_page(soup):
    phones = []
    phone_regex = re.compile(r'\(?\d{2,3}\)?\s?\d{4,5}-?\d{4}')
    
    # Procurar por números de telefone no texto da página
    for text in soup.stripped_strings:
        if phone_regex.match(text):
            phones.append(text)
    
    return phones

# Função para buscar o nome da empresa da página
def get_company_name(soup):
    company_name = None
    title_tag = soup.find('title')
    if title_tag:
        company_name = title_tag.text.strip()

    if not company_name:
        meta_tag = soup.find('meta', {'name': 'author'})
        if meta_tag:
            company_name = meta_tag.get('content', '').strip()

    if not company_name:
        h1_tag = soup.find('h1')
        if h1_tag:
            company_name = h1_tag.text.strip()

    return company_name if company_name else "Nome não encontrado"

# Função para criar e salvar um documento Word com os resultados
def save_to_docx(results):
    doc = Document()
    doc.add_heading('Resultados de Pesquisa', 0)

    for result in results:
        doc.add_heading(f"Empresa: {result['company_name']}", level=1)
        doc.add_paragraph(f"Site: {result['company_url']}")
        doc.add_paragraph(f"Telefones encontrados: {', '.join(result['phones']) if result['phones'] else 'Nenhum telefone encontrado'}")
        doc.add_paragraph(f"E-mails encontrados: {', '.join(result['emails']) if result['emails'] else 'Nenhum e-mail encontrado'}")
        doc.add_parágrafo("\n" + "-"*50 + "\n")

    doc.save("resultados_pesquisa_castanha_para.docx")
    print("Arquivo 'resultados_pesquisa_castanha_para.docx' salvo com sucesso!")

# Função para extrair e-mails de uma página
def get_emails_from_page(soup):
    emails = []
    email_regex = re.compile(r'mailto:[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    
    # Procurar links de e-mails na página
    for link in soup.find_all('a', href=True):
        if email_regex.match(link['href']):
            emails.append(link['href'])
    
    return emails

# Função para realizar a pesquisa no Google e obter links
def search_google(query, driver):
    driver.get("https://www.google.com")
    search_box = driver.find_element(By.NAME, "q")
    search_box.send_keys(query)
    search_box.send_keys(Keys.RETURN)
    time.sleep(3)  # Esperar os resultados carregarem
    
    links = []
    search_results = driver.find_elements(By.CSS_SELECTOR, 'div.yuRUbf > a')
    for result in search_results:
        links.append(result.get_attribute("href"))
    return links

# Função para buscar informações para várias empresas
def extract_data_from_search(query):
    driver = get_driver()

    # Realizar a pesquisa no Google e obter links
    links = search_google(query, driver)

    results = []
    for company_url in links:
        print(f"Buscando dados para: {company_url}")
        
        # Acessar o site
        driver.get(company_url)
        time.sleep(3)  # Aguardar o carregamento da página

        # Tentar encontrar o link de contato e navegar para ele
        try:
            contact_link = None
            # Vamos tentar encontrar o link de contato de diversas maneiras
            possible_contact_links = [
                "contato", "contact", "contatos", "fale conosco", "contato", "contacts"
            ]
            
            for link_text in possible_contact_links:
                try:
                    contact_link = driver.find_element(By.PARTIAL_LINK_TEXT, link_text)
                    break
                except Exception:
                    continue  # Não encontrou, tenta o próximo

            if contact_link:
                contact_link.click()  # Clica no link de contato
                time.sleep(3)  # Aguardar o carregamento da página de contato
            else:
                print(f"Link de contato não encontrado para {company_url}")
                continue  # Caso não tenha encontrado o link de contato, ignora essa página

        except Exception as e:
            print(f"Erro ao tentar acessar a página de contato para {company_url}: {e}")
            continue  # Caso o clique no link de contato falhe, pula para o próximo

        # Extrair o conteúdo da página
        page_source = driver.page_source
        soup = BeautifulSoup(page_source, 'html.parser')

        # Buscar telefones e e-mails na página de contato
        phones = get_phones_from_contact_page(soup)
        emails = get_emails_from_page(soup)

        # Extrair o nome da empresa
        company_name = get_company_name(soup)

        # Adicionar o resultado à lista
        result = {
            'company_name': company_name,
            'company_url': company_url,
            'phones': phones,
            'emails': emails
        }
        results.append(result)

        time.sleep(1)  # Atraso para evitar sobrecarga no servidor

    driver.quit()  # Fechar o driver após processar todos os links

    if results:
        save_to_docx(results)  # Salvando todos os resultados em um arquivo docx
    else:
        print("Nenhum dado encontrado.")

# Lista de empresas para testar
empresas = [
    "Florenzano Nuts",
    "Castanha do Pará Nutriaco",
    "Imaflora",
    "Fábrica de Castanha Benedito Mutran e Cia Ltda",
    # Adicione mais empresas se necessário
]

if __name__ == "__main__":
    for empresa in empresas:
        extract_data_from_search(empresa)
