from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.keys import Keys
import time
import re
from bs4 import BeautifulSoup
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document

# Função para inicializar o Selenium
def get_driver():
    driver_service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=driver_service)
    return driver

# Função para extrair os telefones de uma página
def get_phones_from_contact_page(soup):
    phones = []
    phone_regex = re.compile(r'\(?\d{2,3}\)?\s?\d{4,5}-?\d{4}')
    
    # Procurar por números de telefone no texto da página
    for text in soup.stripped_strings:
        if phone_regex.match(text):
            phones.append(text)
    
    return phones

# Função para buscar o nome da empresa da página
def get_company_name(soup):
    company_name = None
    title_tag = soup.find('title')
    if title_tag:
        company_name = title_tag.text.strip()

    if not company_name:
        meta_tag = soup.find('meta', {'name': 'author'})
        if meta_tag:
            company_name = meta_tag.get('content', '').strip()

    if not company_name:
        h1_tag = soup.find('h1')
        if h1_tag:
            company_name = h1_tag.text.strip()

    return company_name if company_name else "Nome não encontrado"

# Função para criar e salvar um documento Word com os resultados
def save_to_docx(results):
    doc = Document()
    doc.add_heading('Resultados de Pesquisa', 0)

    for result in results:
        doc.add_heading(f"Empresa: {result['company_name']}", level=1)
        doc.add_paragraph(f"Site: {result['company_url']}")
        doc.add_paragraph(f"Telefones encontrados: {', '.join(result['phones']) if result['phones'] else 'Nenhum telefone encontrado'}")
        doc.add_paragraph(f"E-mails encontrados: {', '.join(result['emails']) if result['emails'] else 'Nenhum e-mail encontrado'}")
        doc.add_paragraph("\n" + "-"*50 + "\n")

    doc.save("resultados_pesquisa_castanha_para.docx")
    print("Arquivo 'resultados_pesquisa_castanha_para.docx' salvo com sucesso!")

# Função para extrair e-mails de uma página
def get_emails_from_page(soup):
    emails = []
    email_regex = re.compile(r'mailto:[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    
    # Procurar links de e-mails na página
    for link in soup.find_all('a', href=True):
        if email_regex.match(link['href']):
            emails.append(link['href'])
    
    return emails

# Função principal para realizar a extração de dados
def extract_data_from_search(query):
    # Inicializa o driver
    driver = get_driver()

    # Acessar o site da empresa diretamente
    driver.get(query)
    time.sleep(3)  # Aguarde o carregamento da página

    # Tentar acessar a página de contato clicando no link "contato" ou "contact"
    try:
        # Tentar localizar links com "contato" ou "contact"
        contact_link = driver.find_element(By.PARTIAL_LINK_TEXT, "contato")  # Pode ser 'contact' ou 'contato'
        contact_link.click()
        time.sleep(3)  # Aguarde o carregamento da página de contato
    except Exception as e:
        print(f"Erro ao tentar encontrar o link de contato para {query}: {e}")
        driver.quit()
        return

    # Extrair o conteúdo da página de contato
    page_source = driver.page_source
    soup = BeautifulSoup(page_source, 'html.parser')

    # Buscar telefones e e-mails na página de contato
    phones = get_phones_from_contact_page(soup)
    emails = get_emails_from_page(soup)

    # Extrair o nome da empresa
    company_name = get_company_name(soup)

    # Adicionar o resultado à lista
    result = {
        'company_name': company_name,
        'company_url': query,
        'phones': phones,
        'emails': emails
    }

    # Fechar o driver
    driver.quit()

    return result

# Função para buscar informações para várias empresas
def run_for_companies(companies):
    results = []
    for company in companies:
        print(f"Buscando dados para: {company}")
        result = extract_data_from_search(company)
        if result:
            results.append(result)

    if results:
        save_to_docx(results)  # Salvando todos os resultados em um arquivo docx
    else:
        print("Nenhum dado encontrado.")

# Lista de empresas para testar
empresas = [
    "https://florenzanonuts.com/",  # URL direta para o site da Florenzano Nuts
    # Adicione outras URLs aqui
]

if __name__ == "__main__":
    run_for_companies(empresas)
