from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service  # Necessário para a versão mais recente do Selenium
import time
import re
from bs4 import BeautifulSoup
from docx import Document
from webdriver_manager.chrome import ChromeDriverManager  # Facilita o gerenciamento do chromedriver

# Função para inicializar o Selenium e buscar no Google
def search_google(query):
    # Usando o WebDriver Manager para baixar e gerenciar o chromedriver automaticamente
    driver_service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=driver_service)
    driver.get("https://www.google.com")

    # Encontrar a caixa de pesquisa e enviar a consulta
    search_box = driver.find_element(By.NAME, "q")
    search_box.send_keys(query)
    search_box.send_keys(Keys.RETURN)

    time.sleep(3)  # Aguarde a página carregar

    # Obter links dos resultados de pesquisa
    links = []
    search_results = driver.find_elements(By.CSS_SELECTOR, 'div.yuRUbf > a')
    for result in search_results:
        links.append(result.get_attribute("href"))

    driver.quit()  # Fechar o navegador

    return links

# Função para extrair telefones
def get_phones(soup):
    phones = []
    phone_regex = re.compile(r'\(?\d{2,3}\)?\s?\d{4,5}-?\d{4}')
    for text in soup.stripped_strings:
        if phone_regex.match(text):
            phones.append(text)
    return phones

# Função para extrair e-mails
def get_emails(soup):
    emails = []
    email_regex = re.compile(r'mailto:[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    for link in soup.find_all('a', href=True):
        if email_regex.match(link['href']):
            emails.append(link['href'])
    return emails

# Função para buscar o nome da empresa
def get_company_name(soup):
    company_name = None
    title_tag = soup.find('title')
    if title_tag:
        company_name = title_tag.text.strip()
    if not company_name:
        meta_tag = soup.find('meta', {'name': 'author'})
        if meta_tag:
            company_name = meta_tag.get('content', '').strip()
        if not company_name:
            h1_tag = soup.find('h1')
            if h1_tag:
                company_name = h1_tag.text.strip()
    return company_name if company_name else "Nome não encontrado"

# Função para criar e salvar um documento Word com os resultados
def save_to_docx(results):
    doc = Document()
    doc.add_heading('Resultados de Pesquisa', 0)

    for result in results:
        doc.add_heading(f"Empresa: {result['company_name']}", level=1)
        doc.add_paragraph(f"Site: {result['company_url']}")
        doc.add_paragraph(f"Telefones encontrados: {', '.join(result['phones']) if result['phones'] else 'Nenhum telefone encontrado'}")
        doc.add_paragraph(f"E-mails encontrados: {', '.join(result['emails']) if result['emails'] else 'Nenhum e-mail encontrado'}")
        doc.add_paragraph("\n" + "-"*50 + "\n")

    doc.save("resultados_pesquisa_castanha_para.docx")
    print("Arquivo 'resultados_pesquisa_castanha_para.docx' salvo com sucesso!")

# Função principal que realiza a extração de dados
def extract_data_from_search(query):
    # Buscar resultados no Google usando Selenium
    links = search_google(query)

    results = []
    for company_url in links:
        print(f"Buscando dados para: {query}")
        print(f"Site encontrado: {company_url}")

        # Acessar cada site encontrado e tentar extrair as informações
        driver_service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=driver_service)
        driver.get(company_url)

        time.sleep(3)  # Aguarde o carregamento da página

        page_source = driver.page_source
        driver.quit()  # Fechar o navegador

        # Analisar a página com BeautifulSoup
        company_soup = BeautifulSoup(page_source, 'html.parser')

        # Extrair dados da página
        company_name = get_company_name(company_soup)
        phones = get_phones(company_soup)
        emails = get_emails(company_soup)

        # Adicionar o resultado à lista
        result = {
            'company_name': company_name,
            'company_url': company_url,
            'phones': phones,
            'emails': emails
        }
        results.append(result)

        time.sleep(1)  # Atraso para evitar sobrecarga no servidor

    if results:
        save_to_docx(results)  # Salvando todos os resultados em um arquivo docx
    else:
        print("Nenhum dado encontrado.")

if __name__ == "__main__":
    empresas = [
        "Florenzano Nuts",
        "Castanha do Pará Nutriaco",
        "Imaflora",
        "Fábrica de Castanha Benedito Mutran e Cia Ltda",
        "Cooperativa Comaru",
        "Castanheiras do Pará Ltda.",
        "Castanheiras do Brasil",
        "Associação de Produtores de Castanha do Brasil (APCB)",
        "Cooperativa dos Castanheiros de Cacoal (COOCACOAL)",
        "Fazenda Castanheira",
        "Associação dos Produtores de Castanha do Amapá (APCA)",
        "Cooperativa Agroextrativista do Xingu (CAX)",
        "Cooperativa dos Produtores de Castanha do Maranhão (COOPCAST)",
        "Sementes do Brasil (Sembra)",
        "Fazendas Extrativistas do Norte",
        "Indústria Castanheira Brasil",
        "Fazenda São Luiz",
        "Castanha do Amazonas Ltda.",
        "Cooperativa dos Produtores de Castanha do Xingu (COOPCAX)",
        "Castanheira do Brasil Indústria e Comércio",
        "Cooperativa dos Castanheiros do Pará (COOPCASTANHEIRO)",
        "Agroindústria Castanheiras do Brasil",
        "Santos Castanha",
        "Indústria de Castanha-do-Pará São Francisco",
        "Amazônia Castanha",
        "Cooperativa Agropecuária de Castanheiras de Rondônia",
        "Castanha do Norte Ltda.",
        "Castanha e Cia",
        "Tipico Ceará"
    ]
    
    for empresa in empresas:
        extract_data_from_search(empresa)
