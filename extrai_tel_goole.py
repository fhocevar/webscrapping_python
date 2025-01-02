import requests
from bs4 import BeautifulSoup
from googlesearch import search
from docx import Document  # Importando a biblioteca para manipulação de documentos Word

def buscar_site(nome_empresa):
    # Busca no Google o site da empresa
    query = f"{nome_empresa} site oficial"
    resultados = search(query, num_results=5)  # Faz uma busca com 5 resultados
    
    # Converte o gerador em uma lista
    resultados_lista = list(resultados)
    
    # Retorna o primeiro link da busca ou None se não houver resultados
    if resultados_lista:
        return resultados_lista[0]
    return None

def extrair_dados_do_site(url):
    # Faz uma requisição HTTP para acessar o conteúdo da página
    response = requests.get(url)
    
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Exemplo: tentar pegar o nome da empresa (usando tag 'h1' como exemplo)
        nome_empresa = soup.find('h1').text if soup.find('h1') else 'Nome não encontrado'
        
        # Encontrar telefone (usando padrão de links 'tel:' para telefones)
        telefones = soup.find_all('a', href=True)
        lista_telefones = [tel.get('href') for tel in telefones if 'tel:' in tel.get('href')]
        
        # Encontrar e-mails (usando padrão de links 'mailto:' para e-mails)
        emails = soup.find_all('a', href=True)
        lista_emails = [email.get('href') for email in emails if 'mailto:' in email.get('href')]
        
        return nome_empresa, lista_telefones, lista_emails
    else:
        return None, [], []

# Função para salvar os dados no arquivo docx
def salvar_resultados_docx(resultados):
    # Criando o documento Word
    doc = Document()
    doc.add_heading('Resultados de Pesquisa', 0)

    # Para cada resultado de empresa, adiciona as informações no documento
    for resultado in resultados:
        nome_empresa, site, telefones, emails = resultado
        
        doc.add_heading(f"Empresa: {nome_empresa}", level=1)
        doc.add_paragraph(f"Site: {site}")
        doc.add_paragraph(f"Telefones encontrados: {', '.join(telefones) if telefones else 'Nenhum telefone encontrado'}")
        doc.add_paragraph(f"E-mails encontrados: {', '.join(emails) if emails else 'Nenhum e-mail encontrado'}")
        doc.add_paragraph("\n" + "-"*50 + "\n")
    
    # Salvando o documento
    doc.save("resultados_empresas.docx")
    print("Arquivo 'resultados_empresas.docx' salvo com sucesso!")

# Lista de empresas para pesquisar
empresas = [
    "Florenzano Nuts",
    "Castanha do Pará Nutriaco",
    "Imaflora",
    "Fábrica de Castanha Benedito Mutran e Cia Ltda",
    "Cooperativa Comaru",
    "Castanheiras do Pará Ltda.",
    "Castanheiras do Brasil",
    "Associação de Produtores de Castanha do Brasil (APCB)",
    "Cooperativa dos Castanheiros de Cacoal (COOCACOAL)",
    "Fazenda Castanheira",
    "Associação dos Produtores de Castanha do Amapá (APCA)",
    "Cooperativa Agroextrativista do Xingu (CAX)",
    "Cooperativa dos Produtores de Castanha do Maranhão (COOPCAST)",
    "Sementes do Brasil (Sembra)",
    "Fazendas Extrativistas do Norte",
    "Indústria Castanheira Brasil",
    "Fazenda São Luiz",
    "Castanha do Amazonas Ltda.",
    "Cooperativa dos Produtores de Castanha do Xingu (COOPCAX)",
    "Castanheira do Brasil Indústria e Comércio",
    "Cooperativa dos Castanheiros do Pará (COOPCASTANHEIRO)",
    "Agroindústria Castanheiras do Brasil",
    "Santos Castanha",
    "Indústria de Castanha-do-Pará São Francisco",
    "Amazônia Castanha",
    "Cooperativa Agropecuária de Castanheiras de Rondônia",
    "Castanha do Norte Ltda.",
    "Castanha e Cia",
    "Tipico Ceará"
]

# Lista para armazenar os resultados
resultados_extracao = []

# Para cada empresa na lista, buscar o site e extrair os dados
for empresa in empresas:
    print(f"\nBuscando dados para: {empresa}")
    site = buscar_site(empresa)
    if site:
        print(f"Site encontrado: {site}")
        nome_empresa, telefones, emails = extrair_dados_do_site(site)
        resultados_extracao.append((nome_empresa, site, telefones, emails))
    else:
        print("Não foi possível encontrar o site.")

# Salvar os resultados em um arquivo .docx
salvar_resultados_docx(resultados_extracao)
